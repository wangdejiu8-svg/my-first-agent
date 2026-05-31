import json
from io import BytesIO
from pathlib import Path
from tempfile import NamedTemporaryFile
from tempfile import TemporaryDirectory
from unittest.mock import patch

from django.conf import settings
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import override_settings
from docx import Document
from rest_framework.test import APITestCase

from ai.document_readers import read_document_text
from ai.embeddings import embed_texts_with_backend
from ai.retrieval import (
    ProbeResult,
    QueryProfile,
    RetrievalResult,
    SearchExecution,
    search_current_conversation_knowledge,
)
from .models import Attachment, Conversation, KnowledgeChunk, KnowledgeDocument, Message, RetrievalLog
from .services import AssistantReply, generate_assistant_reply


User = get_user_model()


@override_settings(OPENAI_API_KEY="")
class ChatApiTests(APITestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username="alice",
            email="alice@example.com",
            password="password123",
        )
        self.other_user = User.objects.create_user(
            username="bob",
            email="bob@example.com",
            password="password123",
        )
        self.client.post(
            "/api/auth/login/",
            {"username": "alice", "password": "password123"},
            format="json",
        )

    def test_conversation_list_only_returns_current_users_conversations(self):
        own = Conversation.objects.create(owner=self.user, title="Mine")
        Conversation.objects.create(owner=self.other_user, title="Other")

        response = self.client.get("/api/conversations/")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(len(response.data), 1)
        self.assertEqual(response.data[0]["id"], own.id)

    def test_conversation_list_is_limited_to_twenty_items(self):
        for index in range(25):
            Conversation.objects.create(owner=self.user, title=f"Conv {index}")

        response = self.client.get("/api/conversations/")

        self.assertEqual(response.status_code, 200)
        self.assertEqual(len(response.data), 20)

    def test_get_messages_denies_other_users_conversation(self):
        other = Conversation.objects.create(owner=self.other_user, title="Other")

        response = self.client.get(f"/api/conversations/{other.id}/messages/")

        self.assertEqual(response.status_code, 404)

    def test_delete_conversation_hard_deletes_related_knowledge_and_attachment_file(self):
        with TemporaryDirectory() as media_root, self.settings(MEDIA_ROOT=media_root):
            conversation = Conversation.objects.create(owner=self.user, title="Delete me")
            attachment = Attachment.objects.create(
                owner=self.user,
                conversation=conversation,
                file=SimpleUploadedFile(
                    "note.docx",
                    b"fake-doc-content",
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                ),
                original_name="note.docx",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                size=16,
            )
            knowledge_document = KnowledgeDocument.objects.create(
                owner=self.user,
                conversation=conversation,
                attachment=attachment,
                title="note.docx",
                source_type=KnowledgeDocument.SOURCE_ATTACHMENT,
                index_status=KnowledgeDocument.STATUS_COMPLETED,
                embedding_backend="local-hash",
            )
            KnowledgeChunk.objects.create(
                document=knowledge_document,
                owner=self.user,
                conversation=conversation,
                attachment=attachment,
                chunk_index=0,
                text="chunk",
                token_count=1,
                embedding_json=[1.0, 0.0],
            )
            file_path = Path(attachment.file.path)
            self.assertTrue(file_path.exists())

            response = self.client.delete(f"/api/conversations/{conversation.id}/")

        self.assertEqual(response.status_code, 204)
        self.assertFalse(Conversation.objects.filter(id=conversation.id).exists())
        self.assertFalse(Attachment.objects.filter(id=attachment.id).exists())
        self.assertFalse(KnowledgeDocument.objects.filter(id=knowledge_document.id).exists())
        self.assertFalse(KnowledgeChunk.objects.filter(document_id=knowledge_document.id).exists())
        self.assertFalse(file_path.exists())

    @patch("apps.chat.views.generate_assistant_reply", return_value="你好，我已收到。")
    def test_send_message_creates_conversation_and_messages(self, _mock_reply):
        response = self.client.post(
            "/api/chat/send/",
            {"content": "Hello"},
            format="json",
        )

        self.assertEqual(response.status_code, 201)
        conversation_id = response.data["conversation"]["id"]
        conversation = Conversation.objects.get(id=conversation_id)
        self.assertEqual(conversation.owner, self.user)
        self.assertEqual(conversation.messages.count(), 2)
        self.assertTrue(
            conversation.messages.filter(role=Message.ROLE_USER, content="Hello").exists()
        )
        self.assertTrue(
            conversation.messages.filter(role=Message.ROLE_ASSISTANT).exists()
        )

    @patch("apps.chat.views.generate_assistant_reply", return_value="文件已收到。")
    def test_send_message_binds_uploaded_file_to_user_message(self, _mock_reply):
        upload_response = self.client.post(
            "/api/files/upload/",
            {
                "file": SimpleUploadedFile(
                    "note.pdf",
                    b"%PDF-1.4 mock",
                    content_type="application/pdf",
                )
            },
            format="multipart",
        )
        self.assertEqual(upload_response.status_code, 201)

        response = self.client.post(
            "/api/chat/send/",
            {
                "content": "请看这个文件",
                "attachment_ids": [upload_response.data["id"]],
            },
            format="json",
        )

        self.assertEqual(response.status_code, 201)
        user_message_id = response.data["user_message"]["id"]
        attachment = Attachment.objects.get(id=upload_response.data["id"])
        self.assertEqual(attachment.message_id, user_message_id)
        self.assertEqual(attachment.conversation_id, response.data["conversation"]["id"])
        self.assertEqual(response.data["user_message"]["attachments"][0]["original_name"], "note.pdf")

    def test_upload_rejects_unsupported_file_type(self):
        response = self.client.post(
            "/api/files/upload/",
            {"file": SimpleUploadedFile("note.txt", b"hello", content_type="text/plain")},
            format="multipart",
        )

        self.assertEqual(response.status_code, 400)

    def test_upload_rejects_mixed_files_without_partial_persistence(self):
        response = self.client.post(
            "/api/files/upload/",
            {
                "file": [
                    SimpleUploadedFile(
                        "note.pdf",
                        b"%PDF-1.4 mock",
                        content_type="application/pdf",
                    ),
                    SimpleUploadedFile("note.txt", b"hello", content_type="text/plain"),
                ]
            },
            format="multipart",
        )

        self.assertEqual(response.status_code, 400)
        self.assertFalse(Attachment.objects.filter(owner=self.user).exists())
        self.assertFalse(KnowledgeDocument.objects.filter(owner=self.user).exists())

    def test_upload_rejects_oversized_file(self):
        oversized = SimpleUploadedFile(
            "large.pdf",
            b"a" * (settings.CHAT_MAX_ATTACHMENT_SIZE_BYTES + 1),
            content_type="application/pdf",
        )

        response = self.client.post(
            "/api/files/upload/",
            {"file": oversized},
            format="multipart",
        )

        self.assertEqual(response.status_code, 400)

    def test_send_message_rejects_too_many_attachments(self):
        attachments = [
            Attachment.objects.create(
                owner=self.user,
                file=SimpleUploadedFile(
                    f"note-{index}.pdf",
                    b"%PDF-1.4 mock",
                    content_type="application/pdf",
                ),
                original_name=f"note-{index}.pdf",
                content_type="application/pdf",
                size=12,
            )
            for index in range(settings.CHAT_MAX_ATTACHMENTS_PER_MESSAGE + 1)
        ]

        response = self.client.post(
            "/api/chat/send/",
            {
                "content": "请查看这些文件",
                "attachment_ids": [attachment.id for attachment in attachments],
            },
            format="json",
        )

        self.assertEqual(response.status_code, 400)

    def test_send_message_invalid_attachment_does_not_create_conversation(self):
        response = self.client.post(
            "/api/chat/send/",
            {
                "content": "Hello",
                "attachment_ids": [999999],
            },
            format="json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertEqual(Conversation.objects.filter(owner=self.user).count(), 0)

    def test_send_message_requires_authentication(self):
        self.client.cookies.clear()

        response = self.client.post(
            "/api/chat/send/",
            {"content": "Hello"},
            format="json",
        )

        self.assertEqual(response.status_code, 401)

    def test_docx_reader_extracts_text(self):
        document = Document()
        document.add_paragraph("项目简介")
        document.add_paragraph("这是一个 AI Agent 简历项目。")
        buffer = BytesIO()
        document.save(buffer)

        temp_path = None
        try:
            with NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                temp_path = temp_file.name
                temp_file.write(buffer.getvalue())
                temp_file.flush()
            text = read_document_text(temp_path, "resume.docx")
        finally:
            if temp_path:
                Path(temp_path).unlink(missing_ok=True)

        self.assertIn("项目简介", text)
        self.assertIn("AI Agent 简历项目", text)

    def test_upload_creates_knowledge_document_and_index_status(self):
        upload_response = self.client.post(
            "/api/files/upload/",
            {
                "file": SimpleUploadedFile(
                    "note.pdf",
                    b"%PDF-1.4 mock",
                    content_type="application/pdf",
                )
            },
            format="multipart",
        )

        self.assertEqual(upload_response.status_code, 201)
        attachment = Attachment.objects.get(id=upload_response.data["id"])
        document = KnowledgeDocument.objects.get(attachment=attachment)
        self.assertEqual(document.owner, self.user)
        self.assertEqual(document.title, "note.pdf")
        self.assertEqual(upload_response.data["knowledge_index_status"], document.index_status)
        self.assertEqual(document.index_status, KnowledgeDocument.STATUS_PENDING)

    def test_upload_docx_creates_pending_index_record(self):
        document = Document()
        document.add_paragraph("RAG MVP 文档测试内容。")
        buffer = BytesIO()
        document.save(buffer)

        upload_response = self.client.post(
            "/api/files/upload/",
            {
                "file": SimpleUploadedFile(
                    "note.docx",
                    buffer.getvalue(),
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            format="multipart",
        )

        self.assertEqual(upload_response.status_code, 201)
        attachment = Attachment.objects.get(id=upload_response.data["id"])
        knowledge_document = KnowledgeDocument.objects.get(attachment=attachment)
        self.assertEqual(knowledge_document.index_status, KnowledgeDocument.STATUS_PENDING)
        self.assertFalse(KnowledgeChunk.objects.filter(document=knowledge_document).exists())

    @patch(
        "apps.chat.views.generate_assistant_reply",
        return_value=AssistantReply(
            content="基于检索结果的回答",
            metadata={
                "sources": [
                    {
                        "attachment_id": 1,
                        "attachment_name": "note.pdf",
                        "chunk_count": 1,
                        "snippets": ["项目说明摘要"],
                    }
                ],
                "used_chunks": [
                    {
                        "id": 1,
                        "attachment_id": 1,
                        "attachment_name": "note.pdf",
                        "chunk_index": 0,
                        "text": "项目说明摘要",
                        "score": 0.9,
                    }
                ],
                "is_rag_answer": True,
                "rag_score": 0.9,
                "retrieval_decision": "retrieve",
                "retrieval_reason": "explicit_document_reference",
                "retrieval_router_label": "",
                "retrieval_relatedness": "related",
                "retrieval_probe_score": None,
                "retrieval_judge_used": False,
                "retrieval_decision_version": "v1-fast-probe",
            },
        ),
    )
    def test_send_message_returns_rag_metadata(self, _mock_reply):
        response = self.client.post(
            "/api/chat/send/",
            {"content": "Hello"},
            format="json",
        )

        self.assertEqual(response.status_code, 201)
        assistant_message = response.data["assistant_message"]
        self.assertTrue(assistant_message["is_rag_answer"])
        self.assertEqual(assistant_message["sources"][0]["attachment_name"], "note.pdf")
        self.assertEqual(assistant_message["used_chunks"][0]["chunk_index"], 0)
        self.assertEqual(assistant_message["rag_score"], 0.9)
        self.assertEqual(assistant_message["retrieval_decision"], "retrieve")
        self.assertEqual(assistant_message["retrieval_reason"], "explicit_document_reference")

    def test_send_message_rebinds_orphan_attachment_knowledge_to_new_conversation(self):
        upload_response = self.client.post(
            "/api/files/upload/",
            {
                "file": SimpleUploadedFile(
                    "note.pdf",
                    b"%PDF-1.4 mock",
                    content_type="application/pdf",
                )
            },
            format="multipart",
        )
        self.assertEqual(upload_response.status_code, 201)

        attachment = Attachment.objects.get(id=upload_response.data["id"])
        document = KnowledgeDocument.objects.get(attachment=attachment)
        self.assertIsNone(document.conversation_id)

        with patch("apps.chat.views.generate_assistant_reply", return_value="已绑定"):
            response = self.client.post(
                "/api/chat/send/",
                {
                    "content": "把这个文件绑定到对话",
                    "attachment_ids": [attachment.id],
                },
                format="json",
            )

        self.assertEqual(response.status_code, 201)
        document.refresh_from_db()
        attachment.refresh_from_db()
        self.assertEqual(document.conversation_id, response.data["conversation"]["id"])
        self.assertEqual(attachment.conversation_id, response.data["conversation"]["id"])
        self.assertEqual(
            KnowledgeChunk.objects.filter(attachment=attachment).count(),
            document.chunks.count(),
        )

    def test_generate_assistant_reply_clears_rag_metadata_on_fallback(self):
        conversation = Conversation.objects.create(owner=self.user, title="Fallback")
        retrieval_result = RetrievalResult(
            used_chunks=[
                {
                    "id": 1,
                    "attachment_id": 1,
                    "attachment_name": "note.docx",
                    "chunk_index": 0,
                    "text": "legacy content",
                    "score": 0.99,
                }
            ],
            sources=[
                {
                    "attachment_id": 1,
                    "attachment_name": "note.docx",
                    "chunk_count": 1,
                    "snippets": ["legacy content"],
                }
            ],
        )

        with patch(
            "apps.chat.services.search_current_conversation_knowledge",
            return_value=retrieval_result,
        ), patch(
            "apps.chat.services.invoke_chat_agent",
            side_effect=RuntimeError("backend down"),
        ):
            reply = generate_assistant_reply(conversation, self.user, "Summarize this")

        self.assertEqual(reply.status, "fallback")
        self.assertEqual(
            reply.metadata,
            {
                "sources": [],
                "used_chunks": [],
                "is_rag_answer": False,
                "rag_score": None,
                "retrieval_decision": "skip",
                "retrieval_reason": "fallback",
                "retrieval_router_label": "",
                "retrieval_relatedness": "unrelated",
                "retrieval_probe_score": None,
                "retrieval_judge_used": False,
                "retrieval_decision_version": "fallback",
            },
        )

    def test_generate_assistant_reply_skips_retrieval_for_small_talk(self):
        conversation = Conversation.objects.create(owner=self.user, title="Greeting")

        with patch("apps.chat.services.search_current_conversation_knowledge") as mock_search, patch(
            "apps.chat.services.invoke_chat_agent",
            return_value="你好",
        ) as mock_invoke:
            reply = generate_assistant_reply(conversation, self.user, "你好")

        self.assertEqual(
            reply.metadata,
            {
                "sources": [],
                "used_chunks": [],
                "is_rag_answer": False,
                "rag_score": None,
                "retrieval_decision": "skip",
                "retrieval_reason": "fast_skip_small_talk",
                "retrieval_router_label": "",
                "retrieval_relatedness": "unrelated",
                "retrieval_probe_score": None,
                "retrieval_judge_used": False,
                "retrieval_decision_version": "v1-fast-probe",
            },
        )
        mock_search.assert_not_called()
        self.assertEqual(mock_invoke.call_args.kwargs["retrieval_reference_block"], "")
        log = RetrievalLog.objects.get(conversation=conversation)
        self.assertEqual(log.metadata_json["reason"], "fast_skip_small_talk")
        self.assertEqual(log.metadata_json["action"], "skip")

    def test_generate_assistant_reply_uses_probe_before_retrieval_and_reuses_search_execution(self):
        conversation = Conversation.objects.create(owner=self.user, title="Probe")
        attachment = Attachment.objects.create(
            owner=self.user,
            conversation=conversation,
            file=SimpleUploadedFile("network-lab.pdf", b"%PDF-1.4 mock", content_type="application/pdf"),
            original_name="network-lab.pdf",
            content_type="application/pdf",
            size=12,
        )
        KnowledgeDocument.objects.create(
            owner=self.user,
            conversation=conversation,
            attachment=attachment,
            title="network-lab.pdf",
            source_type=KnowledgeDocument.SOURCE_ATTACHMENT,
            index_status=KnowledgeDocument.STATUS_COMPLETED,
            embedding_backend="local-hash",
        )
        search_execution = SearchExecution(
            query="解释一下 TCP 三次握手",
            top_k=5,
            query_profile=QueryProfile(
                raw="解释一下 TCP 三次握手",
                compact="解释一下TCP三次握手",
                is_short_keyword=False,
            ),
            hits=[],
        )
        probe_result = ProbeResult(
            search_execution=search_execution,
            top_hits=[type("FakeHit", (), {"score": 0.81, "lexical_match": True})()],
        )
        retrieval_result = RetrievalResult(
            used_chunks=[
                {
                    "id": 1,
                    "attachment_id": attachment.id,
                    "attachment_name": attachment.original_name,
                    "chunk_index": 0,
                    "text": "tcp handshake summary",
                    "score": 0.81,
                    "vector_score": 0.63,
                    "lexical_match": True,
                }
            ],
            sources=[
                {
                    "attachment_id": attachment.id,
                    "attachment_name": attachment.original_name,
                    "chunk_count": 1,
                    "snippets": ["tcp handshake summary"],
                }
            ],
        )

        with patch("apps.chat.services.probe_current_conversation_knowledge", return_value=probe_result), patch(
            "apps.chat.services.search_current_conversation_knowledge",
            return_value=retrieval_result,
        ) as mock_search, patch(
            "apps.chat.services.invoke_chat_agent",
            return_value="三次握手是……",
        ):
            reply = generate_assistant_reply(conversation, self.user, "解释一下 TCP 三次握手")

        self.assertEqual(reply.metadata["retrieval_decision"], "retrieve")
        self.assertEqual(reply.metadata["retrieval_reason"], "probe_accept")
        self.assertEqual(reply.metadata["retrieval_probe_score"], 0.81)
        self.assertTrue(reply.metadata["is_rag_answer"])
        self.assertIs(mock_search.call_args.kwargs["search_execution"], search_execution)

    def test_generate_assistant_reply_probe_rejects_low_signal_query(self):
        conversation = Conversation.objects.create(owner=self.user, title="Probe Reject")
        attachment = Attachment.objects.create(
            owner=self.user,
            conversation=conversation,
            file=SimpleUploadedFile("network-lab.pdf", b"%PDF-1.4 mock", content_type="application/pdf"),
            original_name="network-lab.pdf",
            content_type="application/pdf",
            size=12,
        )
        KnowledgeDocument.objects.create(
            owner=self.user,
            conversation=conversation,
            attachment=attachment,
            title="network-lab.pdf",
            source_type=KnowledgeDocument.SOURCE_ATTACHMENT,
            index_status=KnowledgeDocument.STATUS_COMPLETED,
            embedding_backend="local-hash",
        )
        probe_result = ProbeResult(
            search_execution=SearchExecution(
                query="帮我写一个 Python 冒泡排序",
                top_k=5,
                query_profile=QueryProfile(
                    raw="帮我写一个 Python 冒泡排序",
                    compact="帮我写一个Python冒泡排序",
                    is_short_keyword=False,
                ),
                hits=[],
            ),
            top_hits=[type("FakeHit", (), {"score": 0.22, "lexical_match": False})()],
        )

        with patch("apps.chat.services.probe_current_conversation_knowledge", return_value=probe_result), patch(
            "apps.chat.services.search_current_conversation_knowledge",
        ) as mock_search, patch(
            "apps.chat.services.invoke_chat_agent",
            return_value="这里是冒泡排序示例。",
        ):
            reply = generate_assistant_reply(conversation, self.user, "帮我写一个 Python 冒泡排序")

        self.assertEqual(reply.metadata["retrieval_decision"], "skip")
        self.assertEqual(reply.metadata["retrieval_reason"], "probe_reject_low_signal")
        self.assertFalse(reply.metadata["is_rag_answer"])
        mock_search.assert_not_called()
        log = RetrievalLog.objects.get(conversation=conversation)
        self.assertEqual(log.metadata_json["reason"], "probe_reject_low_signal")
        self.assertEqual(log.metadata_json["probe_score"], 0.22)

    def test_generate_assistant_reply_does_not_fail_when_retrieval_log_write_fails(self):
        conversation = Conversation.objects.create(owner=self.user, title="Log Failure")

        with patch(
            "apps.chat.services.log_retrieval_event",
            side_effect=RuntimeError("log write failed"),
        ), patch(
            "apps.chat.services.invoke_chat_agent",
            return_value="你好",
        ):
            reply = generate_assistant_reply(conversation, self.user, "你好")

        self.assertEqual(reply.status, "langgraph_completed")
        self.assertEqual(reply.content, "你好")
        self.assertEqual(reply.metadata["retrieval_decision"], "skip")

    def test_search_current_conversation_knowledge_indexes_legacy_attachment_on_demand(self):
        conversation = Conversation.objects.create(owner=self.user, title="Legacy")
        document = Document()
        document.add_paragraph("legacy retrieval document")
        buffer = BytesIO()
        document.save(buffer)

        attachment = Attachment.objects.create(
            owner=self.user,
            conversation=conversation,
            file=SimpleUploadedFile(
                "legacy.docx",
                buffer.getvalue(),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
            original_name="legacy.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            size=len(buffer.getvalue()),
        )

        self.assertFalse(KnowledgeDocument.objects.filter(attachment=attachment).exists())

        result = search_current_conversation_knowledge(
            user=self.user,
            conversation=conversation,
            query="legacy retrieval",
            top_k=3,
        )

        knowledge_document = KnowledgeDocument.objects.get(attachment=attachment)
        self.assertEqual(knowledge_document.index_status, KnowledgeDocument.STATUS_COMPLETED)
        self.assertTrue(result.is_rag_answer)
        self.assertTrue(any(chunk["attachment_id"] == attachment.id for chunk in result.used_chunks))

    def test_search_current_conversation_knowledge_supports_mixed_embedding_backends(self):
        conversation = Conversation.objects.create(owner=self.user, title="Mixed")
        local_attachment = Attachment.objects.create(
            owner=self.user,
            conversation=conversation,
            file=SimpleUploadedFile("local.pdf", b"%PDF-1.4 mock", content_type="application/pdf"),
            original_name="local.pdf",
            content_type="application/pdf",
            size=12,
        )
        openai_attachment = Attachment.objects.create(
            owner=self.user,
            conversation=conversation,
            file=SimpleUploadedFile("remote.pdf", b"%PDF-1.4 mock", content_type="application/pdf"),
            original_name="remote.pdf",
            content_type="application/pdf",
            size=12,
        )
        local_document = KnowledgeDocument.objects.create(
            owner=self.user,
            conversation=conversation,
            attachment=local_attachment,
            title="local.pdf",
            source_type=KnowledgeDocument.SOURCE_ATTACHMENT,
            index_status=KnowledgeDocument.STATUS_COMPLETED,
            embedding_backend="local-hash",
        )
        openai_document = KnowledgeDocument.objects.create(
            owner=self.user,
            conversation=conversation,
            attachment=openai_attachment,
            title="remote.pdf",
            source_type=KnowledgeDocument.SOURCE_ATTACHMENT,
            index_status=KnowledgeDocument.STATUS_COMPLETED,
            embedding_backend="openai",
        )
        KnowledgeChunk.objects.create(
            document=local_document,
            owner=self.user,
            conversation=conversation,
            attachment=local_attachment,
            chunk_index=0,
            text="local backend chunk",
            token_count=3,
            embedding_json=[1.0, 0.0],
        )
        KnowledgeChunk.objects.create(
            document=openai_document,
            owner=self.user,
            conversation=conversation,
            attachment=openai_attachment,
            chunk_index=0,
            text="openai backend chunk",
            token_count=3,
            embedding_json=[0.0, 1.0, 0.0],
        )

        def fake_embed_texts_for_backend(_texts, backend_name):
            if backend_name == "local-hash":
                return [[1.0, 0.0]]
            if backend_name == "openai":
                return [[0.0, 1.0, 0.0]]
            raise AssertionError(f"unexpected backend {backend_name}")

        with patch("ai.retrieval.embed_texts_for_backend", side_effect=fake_embed_texts_for_backend):
            result = search_current_conversation_knowledge(
                user=self.user,
                conversation=conversation,
                query="backend",
                top_k=5,
            )

        attachment_names = {chunk["attachment_name"] for chunk in result.used_chunks}
        self.assertEqual(attachment_names, {"local.pdf", "remote.pdf"})

    def test_search_current_conversation_knowledge_short_circuits_without_completed_documents(self):
        conversation = Conversation.objects.create(owner=self.user, title="Empty")

        with patch("ai.retrieval.embed_texts_for_backend") as mock_embed:
            result = search_current_conversation_knowledge(
                user=self.user,
                conversation=conversation,
                query="summarize",
                top_k=5,
            )

        self.assertFalse(result.used_chunks)
        self.assertFalse(result.is_rag_answer)
        mock_embed.assert_not_called()

    @override_settings(RAG_MIN_SCORE=0.8)
    def test_search_current_conversation_knowledge_filters_low_score_chunks(self):
        conversation = Conversation.objects.create(owner=self.user, title="Threshold")
        attachment = Attachment.objects.create(
            owner=self.user,
            conversation=conversation,
            file=SimpleUploadedFile("note.pdf", b"%PDF-1.4 mock", content_type="application/pdf"),
            original_name="note.pdf",
            content_type="application/pdf",
            size=12,
        )
        document = KnowledgeDocument.objects.create(
            owner=self.user,
            conversation=conversation,
            attachment=attachment,
            title="note.pdf",
            source_type=KnowledgeDocument.SOURCE_ATTACHMENT,
            index_status=KnowledgeDocument.STATUS_COMPLETED,
            embedding_backend="local-hash",
        )
        KnowledgeChunk.objects.create(
            document=document,
            owner=self.user,
            conversation=conversation,
            attachment=attachment,
            chunk_index=0,
            text="low relevance chunk",
            token_count=3,
            embedding_json=[1.0, 0.0],
        )

        with patch("ai.retrieval.embed_texts_for_backend", return_value=[[0.6, 0.8]]):
            result = search_current_conversation_knowledge(
                user=self.user,
                conversation=conversation,
                query="hello there",
                top_k=5,
            )

        self.assertFalse(result.used_chunks)
        self.assertFalse(result.is_rag_answer)

    def test_search_current_conversation_knowledge_short_keyword_requires_lexical_match(self):
        conversation = Conversation.objects.create(owner=self.user, title="Short Query Gate")
        attachment = Attachment.objects.create(
            owner=self.user,
            conversation=conversation,
            file=SimpleUploadedFile("note.pdf", b"%PDF-1.4 mock", content_type="application/pdf"),
            original_name="note.pdf",
            content_type="application/pdf",
            size=12,
        )
        document = KnowledgeDocument.objects.create(
            owner=self.user,
            conversation=conversation,
            attachment=attachment,
            title="note.pdf",
            source_type=KnowledgeDocument.SOURCE_ATTACHMENT,
            index_status=KnowledgeDocument.STATUS_COMPLETED,
            embedding_backend="local-hash",
        )
        KnowledgeChunk.objects.create(
            document=document,
            owner=self.user,
            conversation=conversation,
            attachment=attachment,
            chunk_index=0,
            text="timer interrupt code example",
            token_count=4,
            embedding_json=[1.0, 0.0],
        )

        with patch("ai.retrieval.embed_texts_for_backend", return_value=[[1.0, 0.0]]):
            result = search_current_conversation_knowledge(
                user=self.user,
                conversation=conversation,
                query="\u7325\u7410",
                top_k=5,
            )

        self.assertFalse(result.used_chunks)
        self.assertFalse(result.sources)
        self.assertFalse(result.is_rag_answer)

    def test_search_current_conversation_knowledge_short_keyword_boosts_exact_match(self):
        conversation = Conversation.objects.create(owner=self.user, title="Short Query Rerank")
        attachment = Attachment.objects.create(
            owner=self.user,
            conversation=conversation,
            file=SimpleUploadedFile("mcu-guide.pdf", b"%PDF-1.4 mock", content_type="application/pdf"),
            original_name="mcu-guide.pdf",
            content_type="application/pdf",
            size=12,
        )
        document = KnowledgeDocument.objects.create(
            owner=self.user,
            conversation=conversation,
            attachment=attachment,
            title="mcu-guide.pdf",
            source_type=KnowledgeDocument.SOURCE_ATTACHMENT,
            index_status=KnowledgeDocument.STATUS_COMPLETED,
            embedding_backend="local-hash",
        )
        KnowledgeChunk.objects.create(
            document=document,
            owner=self.user,
            conversation=conversation,
            attachment=attachment,
            chunk_index=0,
            text="\u5355\u7247\u673a\u5b9e\u9a8c\u4efb\u52a1\u5b9e\u73b0\u6559\u7a0b",
            token_count=10,
            embedding_json=[0.8, 0.6],
        )
        KnowledgeChunk.objects.create(
            document=document,
            owner=self.user,
            conversation=conversation,
            attachment=attachment,
            chunk_index=1,
            text="timer interrupt code example",
            token_count=4,
            embedding_json=[0.95, 0.3122498999],
        )

        with patch("ai.retrieval.embed_texts_for_backend", return_value=[[1.0, 0.0]]):
            result = search_current_conversation_knowledge(
                user=self.user,
                conversation=conversation,
                query="\u5355\u7247\u673a",
                top_k=1,
            )

        self.assertEqual(len(result.used_chunks), 1)
        self.assertEqual(result.used_chunks[0]["chunk_index"], 0)
        self.assertTrue(result.used_chunks[0]["lexical_match"])
        self.assertGreater(result.used_chunks[0]["score"], result.used_chunks[0]["vector_score"])

    def test_search_current_conversation_knowledge_long_query_keeps_vector_only_hit(self):
        conversation = Conversation.objects.create(owner=self.user, title="Long Query Vector")
        attachment = Attachment.objects.create(
            owner=self.user,
            conversation=conversation,
            file=SimpleUploadedFile("arch.pdf", b"%PDF-1.4 mock", content_type="application/pdf"),
            original_name="arch.pdf",
            content_type="application/pdf",
            size=12,
        )
        document = KnowledgeDocument.objects.create(
            owner=self.user,
            conversation=conversation,
            attachment=attachment,
            title="arch.pdf",
            source_type=KnowledgeDocument.SOURCE_ATTACHMENT,
            index_status=KnowledgeDocument.STATUS_COMPLETED,
            embedding_backend="local-hash",
        )
        KnowledgeChunk.objects.create(
            document=document,
            owner=self.user,
            conversation=conversation,
            attachment=attachment,
            chunk_index=0,
            text="timer interrupt code example",
            token_count=4,
            embedding_json=[1.0, 0.0],
        )

        with patch("ai.retrieval.embed_texts_for_backend", return_value=[[1.0, 0.0]]):
            result = search_current_conversation_knowledge(
                user=self.user,
                conversation=conversation,
                query="single chip architecture summary",
                top_k=5,
            )

        self.assertTrue(result.used_chunks)
        self.assertFalse(result.used_chunks[0]["lexical_match"])
        self.assertTrue(result.is_rag_answer)

    @override_settings(
        OPENAI_API_KEY="chat-key",
        OPENAI_BASE_URL="https://chat.example/v1",
        OPENAI_EMBEDDING_API_KEY="embed-key",
        OPENAI_EMBEDDING_BASE_URL="https://embed.example/v1",
        OPENAI_EMBEDDING_MODEL="embed-model",
    )
    def test_embedding_client_uses_dedicated_embedding_configuration(self):
        response = type(
            "EmbeddingResponse",
            (),
            {"data": [type("EmbeddingItem", (), {"embedding": [0.25, 0.75]})()]},
        )()

        with patch("ai.embeddings.OpenAI") as openai_client:
            openai_client.return_value.embeddings.create.return_value = response

            embeddings, backend_name = embed_texts_with_backend(["hello embedding"])

        self.assertEqual(backend_name, "openai")
        self.assertEqual(embeddings, [[0.25, 0.75]])
        openai_client.assert_called_once_with(
            api_key="embed-key",
            base_url="https://embed.example/v1",
        )
        openai_client.return_value.embeddings.create.assert_called_once_with(
            model="embed-model",
            input=["hello embedding"],
        )


@override_settings(OPENAI_API_KEY="")
class ChatStreamApiTests(APITestCase):
    def setUp(self):
        self.user = User.objects.create_user(
            username="stream-user",
            email="stream@example.com",
            password="password123",
        )
        self.client.post(
            "/api/auth/login/",
            {"username": "stream-user", "password": "password123"},
            format="json",
        )

    def test_send_message_stream_returns_incremental_events_and_persists_final_reply(self):
        def fake_stream(_conversation, _user, _user_content):
            yield "Hello"
            yield " world"
            return AssistantReply(
                content="Hello world",
                metadata={
                    "sources": [],
                    "used_chunks": [],
                    "is_rag_answer": False,
                    "rag_score": None,
                    "retrieval_decision": "skip",
                    "retrieval_reason": "fast_skip_small_talk",
                    "retrieval_router_label": "",
                    "retrieval_relatedness": "unrelated",
                    "retrieval_probe_score": None,
                    "retrieval_judge_used": False,
                    "retrieval_decision_version": "v1-fast-probe",
                },
                status="langgraph_completed",
                error="",
            )

        with patch("apps.chat.views.generate_assistant_reply_stream", side_effect=fake_stream):
            response = self.client.post(
                "/api/chat/send-stream/",
                {"content": "Hello"},
                format="json",
            )
            payloads = [
                json.loads(chunk.decode("utf-8"))
                for chunk in response.streaming_content
                if chunk
            ]

        self.assertEqual(response.status_code, 200)
        self.assertTrue(response.streaming)
        self.assertEqual(response["Content-Type"], "application/x-ndjson; charset=utf-8")

        event_types = [item["type"] for item in payloads]
        self.assertEqual(event_types[0], "start")
        self.assertEqual(event_types[-1], "done")
        streamed_text = "".join(item["delta"] for item in payloads if item["type"] == "delta")
        self.assertEqual(streamed_text, "Hello world")
        self.assertEqual(payloads[-1]["assistant_message"]["content"], "Hello world")

        conversation_id = payloads[0]["conversation"]["id"]
        conversation = Conversation.objects.get(id=conversation_id)
        self.assertEqual(conversation.messages.count(), 2)
        self.assertTrue(
            conversation.messages.filter(role=Message.ROLE_ASSISTANT, content="Hello world").exists()
        )
        agent_run = conversation.agent_runs.get()
        self.assertEqual(agent_run.response, "Hello world")
        self.assertEqual(agent_run.status, "langgraph_completed")

    def test_send_message_stream_invalid_attachment_does_not_create_conversation(self):
        response = self.client.post(
            "/api/chat/send-stream/",
            {
                "content": "Hello",
                "attachment_ids": [999999],
            },
            format="json",
        )

        self.assertEqual(response.status_code, 400)
        self.assertEqual(Conversation.objects.filter(owner=self.user).count(), 0)
