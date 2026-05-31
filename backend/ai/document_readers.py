import os
from functools import lru_cache
from pathlib import Path, PurePosixPath
import shutil
import subprocess
import zipfile

from django.conf import settings
from docx import Document
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".docx", ".pdf"}
MAX_DOCUMENT_CHARS = 12000
MAX_DOCX_UNCOMPRESSED_BYTES = 30 * 1024 * 1024
MAX_DOCX_ARCHIVE_ENTRIES = 512
PDF_SIGNATURE = b"%PDF-"
ZIP_SIGNATURES = (b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")
REQUIRED_DOCX_MEMBERS = {
    "[Content_Types].xml",
    "_rels/.rels",
    "word/document.xml",
}
DEFAULT_PDF_OCR_DPI = 300


class DocumentReadError(ValueError):
    """Raised when an uploaded document fails validation or parsing."""


class _LimitedTextBuilder:
    def __init__(self, max_chars, separator):
        self.max_chars = max_chars
        self.separator = separator
        self.parts = []
        self.length = 0
        self.truncated = False

    def add(self, text):
        cleaned = (text or "").strip()
        if not cleaned or self.truncated:
            return

        separator_length = len(self.separator) if self.parts else 0
        remaining = self.max_chars - self.length - separator_length
        if remaining <= 0:
            self.truncated = True
            return

        if len(cleaned) > remaining:
            cleaned = cleaned[:remaining]
            self.truncated = True

        if self.parts:
            self.length += len(self.separator)
        self.parts.append(cleaned)
        self.length += len(cleaned)

    def build(self):
        return self.separator.join(self.parts)


def is_supported_document(filename):
    return Path(filename or "").suffix.lower() in SUPPORTED_EXTENSIONS


def validate_uploaded_document(uploaded_file, *, max_size_bytes):
    original_name = Path(getattr(uploaded_file, "name", "") or "").name
    suffix = Path(original_name).suffix.lower()

    if not original_name:
        raise DocumentReadError("文件名不能为空。")
    if suffix not in SUPPORTED_EXTENSIONS:
        raise DocumentReadError("仅支持上传 .docx 和 .pdf 文件。")

    size = int(getattr(uploaded_file, "size", 0) or 0)
    if size <= 0:
        raise DocumentReadError("不允许上传空文件。")
    if size > max_size_bytes:
        max_size_mb = max_size_bytes // (1024 * 1024)
        raise DocumentReadError(f"单个文件不能超过 {max_size_mb}MB。")

    _validate_document_stream(uploaded_file, suffix)
    return original_name


def read_document_text(file_path, original_name, max_chars=MAX_DOCUMENT_CHARS):
    suffix = Path(original_name or file_path).suffix.lower()
    path = Path(file_path)

    if suffix not in SUPPORTED_EXTENSIONS:
        raise DocumentReadError("当前只支持读取 .docx 和 .pdf 文件。")
    if not path.exists() or not path.is_file():
        raise DocumentReadError("文件不存在或已损坏。")

    with path.open("rb") as document_file:
        _validate_document_stream(document_file, suffix)

    try:
        if suffix == ".docx":
            text, truncated = _read_docx(path, max_chars)
        else:
            text, truncated = _read_pdf(path, max_chars)
    except DocumentReadError:
        raise
    except Exception as exc:
        raise DocumentReadError("文档读取失败，请重新上传后重试。") from exc

    text = _normalize_text(text)
    if not text:
        return ""
    if truncated:
        return f"{text}\n\n[内容过长，已截取前 {max_chars} 个字符。]"
    return text


def _read_docx(path, max_chars):
    document = Document(path)
    builder = _LimitedTextBuilder(max_chars=max_chars, separator="\n")

    for paragraph in document.paragraphs:
        builder.add(paragraph.text)
        if builder.truncated:
            return builder.build(), True

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                builder.add(" | ".join(cells))
            if builder.truncated:
                return builder.build(), True

    return builder.build(), False


def _read_pdf(path, max_chars):
    text, truncated = _read_pdf_with_pypdf(path, max_chars)
    if text:
        return text, truncated

    text, truncated = _read_pdf_with_pymupdf(path, max_chars=max_chars)
    if text:
        return text, truncated

    if not _ocr_fallback_is_configured():
        return "", False

    try:
        ocr_language = _resolve_pdf_ocr_language(
            getattr(settings, "PDF_OCR_LANGUAGE", "auto")
        )
        return _read_pdf_with_pymupdf(
            path,
            max_chars=max_chars,
            use_ocr=True,
            ocr_language=ocr_language,
            ocr_dpi=getattr(settings, "PDF_OCR_DPI", DEFAULT_PDF_OCR_DPI),
        )
    except Exception:
        return "", False


def _read_pdf_with_pypdf(path, max_chars):
    reader = PdfReader(path)
    builder = _LimitedTextBuilder(max_chars=max_chars, separator="\n\n")

    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            builder.add(f"[第 {index} 页]\n{text}")
        if builder.truncated:
            return builder.build(), True

    return builder.build(), False


def _read_pdf_with_pymupdf(
    path,
    *,
    max_chars,
    use_ocr=False,
    ocr_language="eng",
    ocr_dpi=DEFAULT_PDF_OCR_DPI,
):
    try:
        import pymupdf
    except ImportError:
        return "", False

    builder = _LimitedTextBuilder(max_chars=max_chars, separator="\n\n")
    with pymupdf.open(path) as document:
        for index, page in enumerate(document, start=1):
            if use_ocr:
                text = _extract_pdf_page_text_with_ocr(
                    page,
                    ocr_language=ocr_language,
                    ocr_dpi=ocr_dpi,
                )
            else:
                text = page.get_text("text", sort=True).strip()
            if text:
                builder.add(f"[第 {index} 页]\n{text}")
            if builder.truncated:
                return builder.build(), True

    return builder.build(), False


def _extract_pdf_page_text_with_ocr(page, *, ocr_language, ocr_dpi):
    text_page = page.get_textpage_ocr(
        language=ocr_language,
        dpi=ocr_dpi,
        full=True,
    )
    return page.get_text("text", textpage=text_page, sort=True).strip()


def _resolve_pdf_ocr_language(configured_language):
    language = (configured_language or "").strip()
    if language and language.lower() != "auto":
        return language

    available_languages = _get_available_tesseract_languages()
    if {"chi_sim", "eng"}.issubset(available_languages):
        return "chi_sim+eng"
    if "eng" in available_languages:
        return "eng"
    if "chi_sim" in available_languages:
        return "chi_sim"
    return language or "eng"


@lru_cache(maxsize=1)
def _get_available_tesseract_languages():
    if not shutil.which("tesseract"):
        return frozenset()

    try:
        result = subprocess.run(
            ["tesseract", "--list-langs"],
            capture_output=True,
            check=True,
            encoding="utf-8",
            errors="ignore",
        )
    except (OSError, subprocess.SubprocessError):
        return frozenset()

    languages = {
        line.strip()
        for line in result.stdout.splitlines()
        if line.strip() and "available languages" not in line.lower()
    }
    return frozenset(languages)


def _ocr_fallback_is_configured():
    return bool(os.getenv("TESSDATA_PREFIX") or shutil.which("tesseract"))


def _validate_document_stream(file_obj, suffix):
    start_position = None
    if hasattr(file_obj, "tell"):
        try:
            start_position = file_obj.tell()
        except (OSError, ValueError):
            start_position = None

    try:
        if suffix == ".pdf":
            _validate_pdf_stream(file_obj)
        else:
            _validate_docx_stream(file_obj)
    finally:
        if start_position is not None and hasattr(file_obj, "seek"):
            file_obj.seek(start_position)


def _validate_pdf_stream(file_obj):
    if not hasattr(file_obj, "seek"):
        raise DocumentReadError("PDF 文件无法读取。")

    file_obj.seek(0)
    header = file_obj.read(len(PDF_SIGNATURE))
    if header != PDF_SIGNATURE:
        raise DocumentReadError("PDF 文件头无效，请确认上传的是未损坏的 PDF。")


def _validate_docx_stream(file_obj):
    if not hasattr(file_obj, "seek"):
        raise DocumentReadError("DOCX 文件无法读取。")

    file_obj.seek(0)
    header = file_obj.read(4)
    if header not in ZIP_SIGNATURES:
        raise DocumentReadError("DOCX 文件头无效，请确认上传的是标准 Word 文档。")

    file_obj.seek(0)
    try:
        with zipfile.ZipFile(file_obj) as archive:
            members = archive.infolist()
            if not members:
                raise DocumentReadError("DOCX 文件内容为空。")
            if len(members) > MAX_DOCX_ARCHIVE_ENTRIES:
                raise DocumentReadError("DOCX 文件包含异常数量的条目，已拒绝解析。")

            member_names = {info.filename for info in members}
            missing_members = REQUIRED_DOCX_MEMBERS - member_names
            if missing_members:
                raise DocumentReadError("DOCX 文件缺少必要结构，请重新保存后再上传。")

            total_uncompressed_size = 0
            for member in members:
                member_path = PurePosixPath(member.filename)
                if member_path.is_absolute() or ".." in member_path.parts:
                    raise DocumentReadError("DOCX 文件包含非法路径，已拒绝上传。")

                total_uncompressed_size += member.file_size
                if total_uncompressed_size > MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise DocumentReadError("DOCX 解压后体积异常，已拒绝解析。")
    except zipfile.BadZipFile as exc:
        raise DocumentReadError("DOCX 文件已损坏或不是有效的 ZIP 文档。") from exc


def _normalize_text(text):
    lines = [line.strip() for line in (text or "").splitlines()]
    normalized = []
    last_blank = False
    for line in lines:
        if not line:
            if not last_blank:
                normalized.append("")
            last_blank = True
            continue
        normalized.append(line)
        last_blank = False
    return "\n".join(normalized).strip()
